from typing import Dict, List, Any
import PyPDF2
import docx
import io
from .base_text_scraper import BaseTextScraper

class DocumentScraper(BaseTextScraper):
    """Scraper for document files (PDF, DOC, DOCX, etc.)"""
    
    def __init__(self):
        super().__init__(name="document")
        self.supported_formats = ['pdf', 'docx', 'txt']
        
    async def scrape(self, params: Dict[str, Any]) -> List[Dict]:
        """Scrape content from document files"""
        url = params.get('url')
        format_type = params.get('format', '').lower()
        
        if not url:
            self.logger.error("No document URL provided")
            return []
            
        if format_type not in self.supported_formats:
            self.logger.error(f"Unsupported document format: {format_type}")
            return []
            
        try:
            # Download document content
            content = await self._fetch_url(url)
            if not content:
                return []
                
            # Parse document based on format
            document_content = await self._parse_document(content, format_type)
            
            return [{
                'url': url,
                'format': format_type,
                'content': document_content,
                'metadata': await self._extract_metadata(content, format_type)
            }]
            
        except Exception as e:
            self.logger.error(f"Error scraping document: {str(e)}")
            return []
            
    async def _parse_document(self, content: bytes, format_type: str) -> str:
        """Parse document content based on format"""
        if format_type == 'pdf':
            return await self._parse_pdf(content)
        elif format_type == 'docx':
            return await self._parse_docx(content)
        elif format_type == 'txt':
            return content.decode('utf-8')
        return ""
        
    async def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF document"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error parsing PDF: {str(e)}")
            return ""
            
    async def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX document"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            self.logger.error(f"Error parsing DOCX: {str(e)}")
            return ""
            
    async def _extract_metadata(self, content: bytes, format_type: str) -> Dict:
        """Extract metadata from document"""
        metadata = {
            'page_count': 0,
            'author': '',
            'creation_date': '',
            'last_modified': ''
        }
        
        try:
            if format_type == 'pdf':
                pdf_file = io.BytesIO(content)
                reader = PyPDF2.PdfReader(pdf_file)
                metadata['page_count'] = len(reader.pages)
                if reader.metadata:
                    metadata['author'] = reader.metadata.get('/Author', '')
                    metadata['creation_date'] = reader.metadata.get('/CreationDate', '')
                    
            elif format_type == 'docx':
                doc_file = io.BytesIO(content)
                doc = docx.Document(doc_file)
                metadata['page_count'] = len(doc.paragraphs)
                # Extract available docx metadata
                
        except Exception as e:
            self.logger.error(f"Error extracting metadata: {str(e)}")
            
        return metadata
